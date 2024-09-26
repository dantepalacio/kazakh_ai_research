import re
from docx import Document

def extract_text_from_docx(file_path):

    doc = Document(file_path)
    
    lines = []
    
    for para in doc.paragraphs:

        if any(run.font.color and run.font.color.rgb == (255, 0, 0) for run in para.runs):
            continue
        text = para.text.strip()
        if text:  
            lines.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and not any(run.font.color and run.font.color.rgb == (255, 0, 0) for run in cell.paragraphs[0].runs):
                    lines.append(text)
    
    return lines

def replace_text_in_docx(input_file, output_file, translated_texts):
    # Load the original document
    doc = Document(input_file)
    
    # Initialize a counter for translated text lines
    text_counter = 0
    
    # Function to replace text while preserving the formatting
    def replace_runs_in_paragraph(paragraph, new_text):
        # Clear current runs in the paragraph
        for run in paragraph.runs:
            run.clear()
        # Add the new translated text into the paragraph with the same formatting
        paragraph.add_run(new_text)

    # Iterate over all paragraphs in the document
    for para in doc.paragraphs:
        # Skip empty paragraphs or paragraphs that shouldn't be replaced
        if para.text.strip() and text_counter < len(translated_texts):
            replace_runs_in_paragraph(para, translated_texts[text_counter])
            text_counter += 1

    # Iterate over tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Only replace if text exists and we have more translated lines
                if cell.text.strip() and text_counter < len(translated_texts):
                    replace_runs_in_paragraph(cell.paragraphs[0], translated_texts[text_counter])
                    text_counter += 1
    
    # Save the modified document
    doc.save(output_file)


def split_text_overlap(text, max_fragment_length, overlap_length):
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)

    fragments = []
    current_fragment = ""

    for sentence in sentences:
        if len(current_fragment) + len(sentence) <= max_fragment_length:
            current_fragment += sentence
        else:
            if current_fragment:
                fragments.append(current_fragment)

            current_fragment = sentence[:overlap_length]
            overlap = sentence[overlap_length:]
            while len(overlap) > max_fragment_length:
                fragments.append(current_fragment)
                current_fragment = overlap[:overlap_length]
                overlap = overlap[overlap_length:]

    if current_fragment:
        fragments.append(current_fragment)

    return fragments

def split_text(text, max_fragment_length):
    return [text[i:i+max_fragment_length] for i in range(0, len(text), max_fragment_length)]


if __name__=='__main__':
    print(extract_text_from_docx('anketa.docx'))